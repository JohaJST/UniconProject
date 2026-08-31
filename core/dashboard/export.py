"""
core/dashboard/export.py
────────────────────────────
Экспорт результатов Result / SelfResult в Excel (.xlsx) и Word (.docx).

Каждый отчёт содержит:
  - шапку с датой формирования;
  - краткую сводку (кол-во попыток, средний %, и т.п.);
  - детальную таблицу со всеми связанными данными (ФИО, поток/категория,
    предмет/тест, проценты, время);
  - (только Excel) отдельный лист-сводку по ученикам / категориям.

Поддерживается необязательная фильтрация через query-параметры:
  - Result:      ?potok_id=<id>
  - SelfResult:  ?ctg_id=<id>

RBAC и sliding-window таймаут дашборда проверяет DashboardSecurityMiddleware
(все view ниже под префиксом /dashboard/) — свои проверки прав не нужны.
"""
from __future__ import annotations

import datetime
import io
from collections import defaultdict

from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from django.shortcuts import render

from core.models import Potok, Result
from core.models.self import SelfCtg, SelfResult


# ═══════════════════════════════════════════════════════════════════════════
# Общие хелперы
# ═══════════════════════════════════════════════════════════════════════════

def _now_str() -> str:
    return datetime.datetime.now().strftime("%d.%m.%Y %H:%M")


def _file_stamp() -> str:
    return datetime.datetime.now().strftime("%Y-%m-%d_%H-%M")


def _fmt_seconds(total_seconds) -> str:
    """32000 -> '8:53:20', 90 -> '1:30'. Пусто/None -> '—'."""
    if not total_seconds:
        return "—"
    total_seconds = int(total_seconds)
    h, rem = divmod(total_seconds, 3600)
    m, s = divmod(rem, 60)
    return f"{h}:{m:02d}:{s:02d}" if h else f"{m}:{s:02d}"


def _safe_avg(values):
    values = [v for v in values if v is not None]
    return round(sum(values) / len(values), 1) if values else 0


# ═══════════════════════════════════════════════════════════════════════════
# RESULT — сбор данных
# ═══════════════════════════════════════════════════════════════════════════

RESULT_HEADERS = [
    "ID", "Фамилия", "Имя", "Username", "Поток", "Должность", "Компания",
    "Предмет", "Тест", "Верных", "Всего вопросов", "Процент, %",
    "Время прохождения", "Дата прохождения",
]

STUDENT_SUMMARY_HEADERS = ["Ученик", "Поток", "Попыток", "Средний %", "Лучший %"]


def _result_queryset(request):
    qs = (
        Result.objects
        .select_related('user', 'user__potok', 'test', 'test__subject', 'test__potok')
        .order_by('-created')
    )
    potok_id = request.GET.get("potok_id")
    if potok_id and potok_id.isdigit():
        qs = qs.filter(user__potok_id=int(potok_id))
    return qs


def _result_row(r: Result) -> list:
    user = r.user
    test = r.test
    return [
        r.id,
        user.last_name if user else "— (удалён)",
        user.name if user else "—",
        user.username if user else "—",
        user.potok.date_range if (user and user.potok) else "—",
        user.position if (user and user.position) else "—",
        user.company_name if (user and user.company_name) else "—",
        test.subject.name if (test and test.subject) else "—",
        str(test) if test else "— (удалён)",
        r.result if r.result is not None else "—",
        r.totalQuestions if r.totalQuestions is not None else "—",
        r.foyiz if r.foyiz is not None else "—",
        _fmt_seconds(r.time),
        r.created.strftime("%d.%m.%Y %H:%M") if r.created else "—",
    ]


def _result_student_summary(results) -> list:
    groups = defaultdict(list)
    for r in results:
        if r.user:
            key = (
                f"{r.user.last_name} {r.user.name}",
                r.user.potok.date_range if r.user.potok else "—",
            )
        else:
            key = ("Пользователь удалён", "—")
        groups[key].append(r.foyiz or 0)

    rows = []
    for (name, potok), scores in groups.items():
        rows.append([name, potok, len(scores), _safe_avg(scores), max(scores)])
    rows.sort(key=lambda row: row[3], reverse=True)
    return rows


# ═══════════════════════════════════════════════════════════════════════════
# SELF RESULT — сбор данных
# ═══════════════════════════════════════════════════════════════════════════

SELF_RESULT_HEADERS = [
    "ID", "Фамилия", "Имя", "Категория", "Верных", "Всего вопросов",
    "Процент, %", "Первое прохождение", "Дата и время попытки",
]

CTG_SUMMARY_HEADERS = ["Категория", "Попыток", "Уникальных учеников", "Средний %", "Лучший %"]


def _self_result_queryset(request):
    qs = (
        SelfResult.objects
        .select_related('user', 'ctg')
        .order_by('-updated', '-id')
    )
    ctg_id = request.GET.get("ctg_id")
    if ctg_id and ctg_id.isdigit():
        qs = qs.filter(ctg_id=int(ctg_id))
    return qs


def _self_result_row(r: SelfResult) -> list:
    user = r.user
    return [
        r.id,
        user.last_name if user else "— (удалён)",
        user.first_name if user else "—",
        r.ctg.name if r.ctg else "Без категории",
        r.score,
        r.totalQuestions if r.totalQuestions else "—",
        round(r.foiz, 1) if r.foiz is not None else "—",
        user.created.strftime("%d.%m.%Y") if (user and user.created) else "—",
        r.updated.strftime("%d.%m.%Y %H:%M") if r.updated else (
            r.created.strftime("%d.%m.%Y") if r.created else "—"
        ),
    ]


def _self_result_ctg_summary(results) -> list:
    groups = defaultdict(list)
    users_by_ctg = defaultdict(set)
    for r in results:
        name = r.ctg.name if r.ctg else "Без категории"
        groups[name].append(r.foiz or 0)
        if r.user_id:
            users_by_ctg[name].add(r.user_id)

    rows = []
    for name, scores in groups.items():
        rows.append([name, len(scores), len(users_by_ctg[name]), _safe_avg(scores), round(max(scores), 1)])
    rows.sort(key=lambda row: row[3], reverse=True)
    return rows


# ═══════════════════════════════════════════════════════════════════════════
# EXCEL — общий движок листа
# ═══════════════════════════════════════════════════════════════════════════

def _write_excel_sheet(ws, headers, rows, title, summary_lines=None):
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="4F46E5")
    thin = Side(style="thin", color="D1D5DB")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)

    cursor = 1

    ws.merge_cells(start_row=cursor, start_column=1, end_row=cursor, end_column=max(len(headers), 2))
    title_cell = ws.cell(row=cursor, column=1, value=title)
    title_cell.font = Font(bold=True, size=14)
    cursor += 1

    ws.cell(row=cursor, column=1, value=f"Сформировано: {_now_str()}").font = Font(italic=True, size=9, color="6B7280")
    cursor += 2

    if summary_lines:
        for label, value in summary_lines:
            ws.cell(row=cursor, column=1, value=label).font = Font(bold=True)
            ws.cell(row=cursor, column=2, value=value)
            cursor += 1
        cursor += 1

    header_row = cursor
    for col, h in enumerate(headers, start=1):
        cell = ws.cell(row=header_row, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center
        cell.border = border
    cursor += 1

    if not rows:
        ws.cell(row=cursor, column=1, value="Нет данных для отображения").font = Font(italic=True, color="9CA3AF")
        cursor += 1
    else:
        for row_data in rows:
            for col, value in enumerate(row_data, start=1):
                cell = ws.cell(row=cursor, column=col, value=value)
                cell.border = border
                cell.alignment = Alignment(vertical="center", wrap_text=True)
            cursor += 1

    ws.freeze_panes = ws.cell(row=header_row + 1, column=1).coordinate

    for col, h in enumerate(headers, start=1):
        letter = get_column_letter(col)
        max_len = len(str(h))
        for row_data in rows:
            v = row_data[col - 1]
            max_len = max(max_len, len(str(v)) if v is not None else 1)
        ws.column_dimensions[letter].width = min(max(max_len + 2, 10), 48)

    return cursor


def _xlsx_response(wb, filename: str) -> HttpResponse:
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    response = HttpResponse(
        buf.read(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


# ═══════════════════════════════════════════════════════════════════════════
# WORD — общий движок документа
# ═══════════════════════════════════════════════════════════════════════════

def _shade_cell(cell, color_hex: str):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def _repeat_header_row(row):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement('w:tblHeader')
    header.set(qn('w:val'), "true")
    tr_pr.append(header)


def _add_table(doc, headers, rows, header_color="4F46E5"):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade_cell(hdr_cells[i], header_color)
    _repeat_header_row(table.rows[0])

    if not rows:
        row_cells = table.add_row().cells
        row_cells[0].text = "Нет данных для отображения"
    else:
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, value in enumerate(row_data):
                row_cells[i].text = "" if value is None else str(value)
                for p in row_cells[i].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.size = Pt(8.5)

    return table


def _build_docx_report(title, headers, rows, summary_lines, summary_title, summary_headers, summary_rows):
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT

    doc = Document()

    # Альбомная ориентация — иначе широкая таблица не влезает.
    section = doc.sections[0]
    new_w, new_h = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = new_w, new_h
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)

    doc.add_heading(title, level=1)

    meta = doc.add_paragraph()
    meta_run = meta.add_run(f"Сформировано: {_now_str()}")
    meta_run.italic = True
    meta_run.font.size = Pt(9)

    if summary_lines:
        p = doc.add_paragraph()
        for label, value in summary_lines:
            p.add_run(f"{label} ").bold = True
            p.add_run(f"{value}    ")

    doc.add_paragraph()
    doc.add_heading("Детализация", level=2)
    _add_table(doc, headers, rows)

    if summary_rows:
        doc.add_paragraph()
        doc.add_heading(summary_title, level=2)
        _add_table(doc, summary_headers, summary_rows, header_color="0F172A")

    return doc


def _docx_response(doc, filename: str) -> HttpResponse:
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    response = HttpResponse(
        buf.read(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


# ═══════════════════════════════════════════════════════════════════════════
# VIEWS — Result
# ═══════════════════════════════════════════════════════════════════════════

@login_required(login_url="login")
def export_result_excel(request):
    from openpyxl import Workbook

    results = list(_result_queryset(request))
    rows = [_result_row(r) for r in results]

    total = len(results)
    scores = [r.foyiz for r in results if r.foyiz is not None]
    summary = [
        ("Всего попыток:", total),
        ("Средний процент:", f"{_safe_avg(scores)}%"),
        ("Отлично (≥75%):", sum(1 for s in scores if s >= 75)),
        ("Слабо (<50%):", sum(1 for s in scores if s < 50)),
    ]

    wb = Workbook()
    ws = wb.active
    ws.title = "Результаты тестов"
    _write_excel_sheet(ws, RESULT_HEADERS, rows, "Отчёт по результатам тестов", summary)

    ws2 = wb.create_sheet("Сводка по ученикам")
    _write_excel_sheet(ws2, STUDENT_SUMMARY_HEADERS, _result_student_summary(results), "Сводка по ученикам")

    return _xlsx_response(wb, f"results_{_file_stamp()}.xlsx")


@login_required(login_url="login")
def export_result_word(request):
    results = list(_result_queryset(request))
    rows = [_result_row(r) for r in results]

    total = len(results)
    scores = [r.foyiz for r in results if r.foyiz is not None]
    summary = [
        ("Всего попыток:", total),
        ("Средний процент:", f"{_safe_avg(scores)}%"),
        ("Отлично (≥75%):", sum(1 for s in scores if s >= 75)),
    ]

    doc = _build_docx_report(
        title="Отчёт по результатам тестов",
        headers=RESULT_HEADERS,
        rows=rows,
        summary_lines=summary,
        summary_title="Сводка по ученикам",
        summary_headers=STUDENT_SUMMARY_HEADERS,
        summary_rows=_result_student_summary(results),
    )
    return _docx_response(doc, f"results_{_file_stamp()}.docx")


# ═══════════════════════════════════════════════════════════════════════════
# VIEWS — SelfResult
# ═══════════════════════════════════════════════════════════════════════════

@login_required(login_url="login")
def export_selfresult_excel(request):
    from openpyxl import Workbook

    results = list(_self_result_queryset(request))
    rows = [_self_result_row(r) for r in results]

    total = len(results)
    scores = [r.foiz for r in results if r.foiz is not None]
    summary = [
        ("Всего попыток:", total),
        ("Средний процент:", f"{_safe_avg(scores)}%"),
        ("Отлично (≥75%):", sum(1 for s in scores if s >= 75)),
    ]

    wb = Workbook()
    ws = wb.active
    ws.title = "Self Check"
    _write_excel_sheet(ws, SELF_RESULT_HEADERS, rows, "Отчёт по Self Check", summary)

    ws2 = wb.create_sheet("Сводка по категориям")
    _write_excel_sheet(ws2, CTG_SUMMARY_HEADERS, _self_result_ctg_summary(results), "Сводка по категориям")

    return _xlsx_response(wb, f"self_check_{_file_stamp()}.xlsx")


@login_required(login_url="login")
def export_selfresult_word(request):
    results = list(_self_result_queryset(request))
    rows = [_self_result_row(r) for r in results]

    total = len(results)
    scores = [r.foiz for r in results if r.foiz is not None]
    summary = [
        ("Всего попыток:", total),
        ("Средний процент:", f"{_safe_avg(scores)}%"),
    ]

    doc = _build_docx_report(
        title="Отчёт по Self Check",
        headers=SELF_RESULT_HEADERS,
        rows=rows,
        summary_lines=summary,
        summary_title="Сводка по категориям",
        summary_headers=CTG_SUMMARY_HEADERS,
        summary_rows=_self_result_ctg_summary(results),
    )
    return _docx_response(doc, f"self_check_{_file_stamp()}.docx")


# ═══════════════════════════════════════════════════════════════════════════
# HUB — страница выбора отчёта
# ═══════════════════════════════════════════════════════════════════════════

@login_required(login_url="login")
def export_hub(request):
    return render(request, "pages/dashboard/export.html", {
        "potoks": Potok.objects.all().order_by('-start'),
        "ctgs": SelfCtg.objects.all().order_by('name_uz'),
    })